"""
AS Code — Document Parser
Soporta: TXT, PDF, DOCX
Objetivo: extraer texto limpio → inyectar en chat
"""

import logging
import os
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Documento parseado listo para chat."""
    filename: str
    file_type: str  # "txt", "pdf", "docx"
    text: str
    char_count: int
    chunk_count: int = 0  # para RAG después


def parse_txt(file_path: str) -> str:
    """Parse .txt — trivial."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # Windows ANSI fallback
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()


def parse_pdf(file_path: str) -> str:
    """Parse .pdf con pypdf (liviano)."""
    try:
        import pypdf
    except ImportError:
        logger.error("pypdf no instalado. Instala con: pip install pypdf")
        raise

    try:
        text = ""
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error parseando PDF: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """Parse .docx con python-docx (liviano)."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx no instalado. Instala con: pip install python-docx")
        raise

    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"Error parseando DOCX: {e}")
        raise


def parse_xlsx(file_path: str) -> str:
    """Parse .xlsx / .xlsm con openpyxl (preservando formato tabular Markdown)."""
    try:
        import openpyxl
    except ImportError:
        logger.error("openpyxl no instalado. Instala con: pip install openpyxl")
        raise

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        MAX_ROWS = 500
        MAX_COLS = 50
        output_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows_data = []
            truncated = False
            row_idx = 0

            for row in sheet.iter_rows(values_only=True):
                if row_idx >= MAX_ROWS:
                    truncated = True
                    break

                # Check if row has any non-None value
                if any(v is not None for v in row):
                    # Limit columns to MAX_COLS
                    row_val = row[:MAX_COLS]
                    row_str_list = []
                    for v in row_val:
                        if v is None:
                            row_str_list.append("")
                        elif isinstance(v, float):
                            row_str_list.append(f"{v:.4f}".rstrip('0').rstrip('.'))
                        else:
                            row_str_list.append(str(v).strip().replace("\n", " ").replace("|", "\\|"))
                    rows_data.append(row_str_list)
                    row_idx += 1

            if not rows_data:
                continue

            if truncated:
                logger.warning(
                    f"Hoja '{sheet_name}' truncada a {MAX_ROWS} filas por límite."
                )

            # Build markdown table
            sheet_md = []
            sheet_md.append(f"### Sheet: {sheet_name}")
            
            headers = rows_data[0]
            headers_cleaned = [h if h != "" else f"Col{i+1}" for i, h in enumerate(headers)]
            
            sheet_md.append("| " + " | ".join(headers_cleaned) + " |")
            sheet_md.append("| " + " | ".join(["---"] * len(headers_cleaned)) + " |")
            
            for row in rows_data[1:]:
                if len(row) < len(headers_cleaned):
                    row += [""] * (len(headers_cleaned) - len(row))
                sheet_md.append("| " + " | ".join(row) + " |")
            
            if truncated:
                sheet_md.append(f"\n*Nota: Hoja '{sheet_name}' truncada a {MAX_ROWS} filas.*")
                
            output_parts.append("\n".join(sheet_md))

        wb.close()
        return "\n\n".join(output_parts).strip()
    except Exception as e:
        logger.error(f"Error parseando XLSX: {e}")
        raise


def parse_document(file_path: str) -> ParsedDocument:
    """
    Router: detecta tipo y parsea.

    Args:
        file_path: ruta al archivo

    Returns:
        ParsedDocument listo para inyectar en chat
    """
    path = Path(file_path)
    suffix = path.suffix.lower()
    filename = path.name

    logger.info(f"Parseando: {filename}")

    if suffix == ".txt":
        text = parse_txt(file_path)
        file_type = "txt"
    elif suffix == ".pdf":
        text = parse_pdf(file_path)
        file_type = "pdf"
    elif suffix in [".docx", ".doc"]:
        text = parse_docx(file_path)
        file_type = "docx"
    elif suffix in [".xlsx", ".xlsm"]:
        text = parse_xlsx(file_path)
        file_type = "xlsx"
    else:
        raise ValueError(f"Formato no soportado: {suffix}")

    doc = ParsedDocument(
        filename=filename,
        file_type=file_type,
        text=text,
        char_count=len(text),
    )

    logger.info(f"✓ Parseado: {filename} ({file_type}) — {doc.char_count} chars")
    return doc


def clean_text(text: str) -> str:
    """
    Limpia texto:
    - Espacios dobles
    - Newlines múltiples
    - Control chars
    """
    text = " ".join(text.split())  # espacios dobles → simple
    text = "\n".join([line.strip() for line in text.split("\n")])
    return text.strip()


def truncate_context(text: str, max_chars: int = 8000) -> str:
    """
    Si el documento es muy largo, trunca.
    Útil para no overflow de contexto en LLM.

    Returns:
        texto truncado + "[... documento truncado]"
    """
    if len(text) > max_chars:
        logger.warning(
            f"Documento truncado de {len(text)} a {max_chars} chars"
        )
        return text[:max_chars] + "\n\n[... documento truncado por tamaño]"
    return text
